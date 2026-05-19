"""Generate Content Revision Checklist Word document."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], "E8E8E8")
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = val
    doc.add_paragraph()


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Content Revision Checklist", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "From annotated one-pager printouts. Mark Keep / Change / Skip as you review."
    )
    doc.add_paragraph(
        "Your notes map primarily to the Visual overview one-pager and shared data in "
        "criteria.ts and onePagerContent.ts. Most copy also appears on the public homepage "
        "(About, WhyUs, InvestmentCriteria) unless you want one-pager-only copy later."
    )

    # 1. Visual overview
    doc.add_heading("1. Visual overview (hero)", level=1)
    add_table(
        doc,
        ["#", "Status", "Current", "Notes from printout"],
        [
            ["1.1", "No change", "We partner with founders to build enduring businesses.", "No annotations"],
            ["1.2", "No change", "Permanent capital. Operational freedom. Long-term value.", "No annotations"],
        ],
    )

    # 2. Who we are
    doc.add_heading("2. Who we are", level=1)
    doc.add_paragraph(
        "Source: ABOUT_LEAD in onePagerContent.ts and duplicate paragraph in About.tsx."
    )
    doc.add_paragraph("Current paragraph:", style="List Bullet")
    doc.add_paragraph(
        "Charlton Bleecker Group LLC is a private holding company focused on acquiring "
        "and growing enduring scalable B2B businesses. We partner with owner-managers to "
        "position companies for long-term success, without fund cycle pressure, LP timelines, "
        "or re-trade risk. Unlike traditional PE, we hold permanently. Your legacy, your team, "
        "and your culture remain intact."
    )
    add_table(
        doc,
        ["#", "Action", "Proposed revision"],
        [
            ["2.1", "Replace", "Change \"without fund cycle pressure\" → \"without the constraints of a fund termination date\""],
            ["2.2", "Delete", "Remove \"LP timelines,\""],
            ["2.3", "Delete", "Remove \"or re-trade risk\""],
            ["2.4", "Replace", "Change \"Unlike traditional PE, we hold permanently.\" → \"Unlike traditional PE, we do not prioritize the exit.\" (Note on printout was cut off — confirm full sentence if different.)"],
            ["2.5", "Add (new)", "Add: \"We welcome challenges, whether operational, financial or succession.\" (Placement TBD: end of paragraph vs. under \"What you get\")"],
            ["2.6", "No change", "Headline \"Built differently. Invested permanently.\""],
        ],
    )
    doc.add_paragraph("After edits (draft for confirmation):", style="Intense Quote")
    doc.add_paragraph(
        "…for long-term success, without the constraints of a fund termination date. "
        "Unlike traditional PE, we do not prioritize the exit. Your legacy, your team, "
        "and your culture remain intact. We welcome challenges, whether operational, "
        "financial or succession."
    )

    # 3. What you get
    doc.add_heading("3. What you get (pillar cards)", level=1)
    doc.add_paragraph("Source: PILLARS in criteria.ts")

    doc.add_heading("3a. Operational Freedom", level=2)
    add_table(doc, ["#", "Action", "Detail"], [
        ["3.1", "Delete", "Remove ending phrase: \"not a new boss.\""],
        ["3.2", "Review", "Red mark next to title — unclear if rename/reorder only"],
    ])
    doc.add_paragraph("Proposed body:", style="Intense Quote")
    doc.add_paragraph(
        "Our model gives managers the freedom to lead, with the backing of experienced "
        "investors focused on long-term value creation. We provide capital and strategic support."
    )

    doc.add_heading("3b. Collaborative Expertise", level=2)
    add_table(doc, ["#", "Action", "Detail"], [
        ["3.3", "Delete", "Remove \"and turnarounds\" from expertise list"],
        ["3.4", "Review", "Red mark next to title"],
    ])
    doc.add_paragraph("Proposed body:", style="Intense Quote")
    doc.add_paragraph(
        "Our team brings deep expertise spanning corporate finance, M&A, capital raising, "
        "governance, and legal, providing both strategic insight and tactical execution from day one."
    )

    doc.add_heading("3c. Sector Focus", level=2)
    add_table(doc, ["#", "Action", "Detail"], [
        ["3.5", "Insert", "Add \"B2B\" before \"industries\""],
        ["3.6", "Review", "Red \"5\" in corner — may indicate reorder; confirm if order should change"],
    ])
    doc.add_paragraph("Proposed opening:", style="Intense Quote")
    doc.add_paragraph("We invest in B2B industries where we can add meaningful value: …")

    # 4. Why founders choose us
    doc.add_heading("4. Why founders choose us (WHY_US cards)", level=1)
    doc.add_paragraph("Source: criteria.ts. Homepage uses same data with different section title.")

    doc.add_heading("4a. Your Legacy, Protected", level=2)
    add_table(doc, ["#", "Action", "Detail"], [
        ["4.1", "Keep", "We don't flip. We hold. (green highlight — no strike-through)"],
        ["4.2", "Delete", "Remove \"with no integration into a portfolio,\" (note: UNNECESSARY)"],
    ])
    doc.add_paragraph("Proposed body:", style="Intense Quote")
    doc.add_paragraph(
        "We don't flip. We hold. Your brand, culture, and people remain intact after the "
        "transaction, no culture reset."
    )

    doc.add_heading("4b. Operational Autonomy", level=2)
    add_table(doc, ["#", "Action", "Detail"], [
        ["4.3", "Delete", "Remove opening: \"Management stays in place.\""],
        ["4.4", "Delete", "Remove ending: \"not a new boss or a corporate playbook.\""],
        ["4.5", "Add", "After \"strategic support,\" insert \"and clear milestones.\""],
        ["4.6", "Add", "Incorporate: \"We give managers the opportunity to do what they do best.\""],
        ["4.7", "Review", "Red arrow to Operational Freedom — confirm merge vs. align wording only"],
    ])
    doc.add_paragraph("Proposed body (draft):", style="Intense Quote")
    doc.add_paragraph(
        "We provide capital and strategic support, and clear milestones. "
        "We give managers the opportunity to do what they do best."
    )

    doc.add_heading("4c. Certainty of Close", level=2)
    add_table(doc, ["#", "Action", "Detail"], [
        ["4.8", "Delete", "Remove \"no fund cycles,\" \"no LP approval committees,\" \"no re-trades.\""],
        ["4.9", "Add", "Emphasize \"No surprises.\""],
        ["4.10", "Review", "Red \"4\" in corner — possible reorder; confirm"],
    ])
    doc.add_paragraph("Proposed body:", style="Intense Quote")
    doc.add_paragraph("Permanent capital means when we say yes, we mean it. No surprises.")

    doc.add_heading("4d. Fair, Transparent Valuations", level=2)
    add_table(doc, ["#", "Action", "Detail"], [
        ["4.11", "Replace", "\"explain our math\" → \"explain our reasoning\""],
        ["4.12", "Insert", "\"after diligence\" → \"after due diligence\""],
    ])
    doc.add_paragraph("Proposed body:", style="Intense Quote")
    doc.add_paragraph(
        "We lead with our number and explain our reasoning. "
        "No games, no bait-and-switch after due diligence."
    )

    doc.add_heading("4e. Speed & Simplicity", level=2)
    add_table(doc, ["#", "Action", "Detail"], [
        ["4.13", "Keep (STET)", "\"Intro call to LOI in weeks, not months.\" — replacement crossed out; (STET.) = leave original"],
    ])

    doc.add_heading("4f. Confidentiality First", level=2)
    add_table(doc, ["#", "Action", "Detail"], [["4.14", "No change", "No annotations"]])

    doc.add_heading("4g. Section-level", level=2)
    add_table(doc, ["#", "Action", "Detail"], [
        ["4.15", "Review", "Large red scribble between card rows — confirm reorder, rename, or ignore"],
    ])

    # 5. Criteria
    doc.add_heading("5. Criteria at a glance", level=1)
    doc.add_paragraph("Source: CRITERIA in criteria.ts. Homepage uses same values under Our Investment Criteria.")
    add_table(
        doc,
        ["#", "Field", "Current", "Proposed"],
        [
            ["5.1", "Annual Revenue", "$3M – $30M", "$3M+ (remove upper cap)"],
            ["5.2", "EBITDA", "$1M – $8M", "$1M+ (remove upper cap)"],
            ["5.3", "EBITDA Margin", "15%+", "15%+ potential (confirm wording)"],
            ["5.4", "Geography", "United States", "No change"],
            ["5.5", "Ownership", "Founder-owned", "No change"],
            ["5.6", "Business Type", "B2B Focus", "No change"],
            ["5.7", "Hold Period", "Permanent", "Long term to permanent (confirm exact phrasing)"],
            ["5.8", "Deal Structure", "Flexible", "Flexible, creative"],
            ["5.9", "Layout (defer)", "Grid on overview one-pager", "USE OTHER FORMAT ON LATER PAGE — defer layout change"],
        ],
    )

    # 6. How it works
    doc.add_heading("6. How it works", level=1)
    add_table(doc, ["#", "Status", "Detail"], [
        ["6.1", "No copy edits", "Visible portion had no marks on printout"],
        ["6.2", "Implicit", "PROCESS_STEPS unchanged unless separate notes for steps 2–6"],
    ])

    # 7. Path to partnership
    doc.add_heading("7. Path to partnership (model one-pager)", level=1)
    add_table(doc, ["#", "Status", "Detail"], [
        ["7.1", "Approved", "Path to partnership heading boxed with YES — keep as-is on model one-pager"],
    ])

    # 8. Ripple effects
    doc.add_heading("8. Ripple effects (after you confirm)", level=1)
    doc.add_paragraph("Implementation would likely touch:", style="List Bullet")
    for item in [
        "src/data/criteria.ts — pillars, criteria, why-us cards",
        "src/data/onePagerContent.ts — ABOUT_LEAD and related strings",
        "src/components/sections/About.tsx — sync ABOUT_LEAD or import to avoid drift",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph(
        "Out of scope unless requested: criteria layout on overview (5.9), card reordering (3.6, 4.10), "
        "merging Operational Freedom + Operational Autonomy (4.7)."
    )

    # Open items
    doc.add_heading("Open items for your reply", level=1)
    items = [
        "2.4 — Full sentence for \"we do not prioritize the exit…\"",
        "2.5 — Where the \"welcome challenges\" line belongs",
        "4.7 — Merge Operational Autonomy and Operational Freedom, or only de-duplicate?",
        "4.15 — Meaning of scribble between card rows",
        "5.7 — Exact Hold Period label text",
        "Scope — Apply to homepage + one-pagers, or overview only?",
    ]
    for i, item in enumerate(items, 1):
        doc.add_paragraph(f"{i}. {item}", style="List Number")

    out_path = (
        r"c:\mSquare Solutions\GitHub\clients\charlton-bleecker"
        r"\docs\Charlton-Bleecker-Content-Revision-Checklist.docx"
    )
    doc.save(out_path)
    print(out_path)


if __name__ == "__main__":
    main()
